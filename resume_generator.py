from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import psycopg2
from psycopg2.extras import RealDictCursor
from database import get_connection
import io

def format_date(date_obj):
    """Format date object to string"""
    if date_obj:
        return date_obj.strftime("%B %Y")
    return "Present"

def add_section_heading(document, text):
    """Add a section heading with underline like in the template"""
    para = document.add_paragraph()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Add bottom border (underline for the section)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    para.paragraph_format.space_after = Pt(6)
    return para

def add_entry_with_date(document, left_text, right_text, bold_left=True):
    """Add an entry with text on left and date on right"""
    para = document.add_paragraph()

    # Add tab stop at right margin for date alignment
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    left_run = para.add_run(left_text)
    left_run.font.name = 'Times New Roman'
    left_run.font.size = Pt(11)
    if bold_left:
        left_run.font.bold = True

    para.add_run('\t')

    right_run = para.add_run(right_text)
    right_run.font.name = 'Times New Roman'
    right_run.font.size = Pt(11)

    para.paragraph_format.space_after = Pt(0)
    return para

def add_subtitle(document, text):
    """Add a subtitle line (like degree or job title)"""
    para = document.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(0)
    return para

def add_bullet_point(document, text):
    """Add a bullet point item"""
    para = document.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(2)
    return para

def generate_resume(user_id: int) -> io.BytesIO:
    """
    Generate a DOCX resume for a user
    Returns BytesIO object containing the document
    """

    # Fetch all user data
    conn = get_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)

    try:
        # Get user info
        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()

        if not user:
            raise ValueError("User not found")

        # Get work experiences
        cursor.execute("""
            SELECT * FROM work_experiences
            WHERE user_id = %s
            ORDER BY is_current DESC, start_date DESC
        """, (user_id,))
        work_experiences = cursor.fetchall()

        # Get education
        cursor.execute("""
            SELECT * FROM education
            WHERE user_id = %s
            ORDER BY end_date DESC NULLS FIRST
        """, (user_id,))
        education = cursor.fetchall()

        # Get skills
        cursor.execute("""
            SELECT * FROM skills
            WHERE user_id = %s
            ORDER BY skill_name ASC
        """, (user_id,))
        skills = cursor.fetchall()

        # Get projects
        cursor.execute("""
            SELECT * FROM projects
            WHERE user_id = %s
            ORDER BY start_date DESC NULLS LAST
        """, (user_id,))
        projects = cursor.fetchall()

    finally:
        cursor.close()
        conn.close()

    # Create document
    document = Document()

    # Set margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # ==================== HEADER - NAME ====================
    name_para = document.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(user['name'])
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.name = 'Times New Roman'
    name_para.paragraph_format.space_after = Pt(0)

    # ==================== CONTACT INFO ====================
    contact_para = document.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = []

    if user.get('location'):
        contact_parts.append(user['location'])
    if user.get('phone'):
        contact_parts.append(user['phone'])
    if user.get('email'):
        contact_parts.append(user['email'])

    contact_run = contact_para.add_run(' | '.join(contact_parts))
    contact_run.font.size = Pt(10)
    contact_run.font.name = 'Times New Roman'
    contact_para.paragraph_format.space_after = Pt(0)

    # Add LinkedIn/URLs if available
    urls_parts = []
    if user.get('linkedin_url'):
        urls_parts.append(user['linkedin_url'])
    if user.get('github_url'):
        urls_parts.append(user['github_url'])

    if urls_parts:
        urls_para = document.add_paragraph()
        urls_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        urls_run = urls_para.add_run(' | '.join(urls_parts))
        urls_run.font.size = Pt(10)
        urls_run.font.name = 'Times New Roman'
        urls_run.font.color.rgb = RGBColor(0, 0, 255)
        urls_para.paragraph_format.space_after = Pt(6)
    else:
        contact_para.paragraph_format.space_after = Pt(6)

    # ==================== EDUCATION ====================
    if education:
        add_section_heading(document, 'EDUCATION')

        for edu in education:
            # School name and location with date on right
            school_text = f"{edu['school']}"
            date_text = format_date(edu['end_date']) if edu.get('end_date') else ""
            add_entry_with_date(document, school_text, date_text, bold_left=True)

            # Degree and field of study
            degree_text = edu['degree']
            if edu.get('field_of_study'):
                degree_text += f", {edu['field_of_study']}"
            add_subtitle(document, degree_text)

            # GPA if available
            if edu.get('gpa'):
                gpa_para = document.add_paragraph()
                run = gpa_para.add_run(f"GPA: {edu['gpa']}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                gpa_para.paragraph_format.space_after = Pt(6)
            else:
                # Add spacing after education entry
                document.add_paragraph().paragraph_format.space_after = Pt(0)

    # ==================== SKILLS ====================
    if skills:
        add_section_heading(document, 'SKILLS')

        # Group skills by proficiency
        skills_by_prof = {}
        for skill in skills:
            prof = skill.get('proficiency') or 'Other'
            if prof not in skills_by_prof:
                skills_by_prof[prof] = []
            skills_by_prof[prof].append(skill['skill_name'])

        # Display skills grouped by proficiency
        for proficiency, skill_list in skills_by_prof.items():
            para = document.add_paragraph()
            prof_run = para.add_run(f"{proficiency}: ")
            prof_run.font.bold = True
            prof_run.font.name = 'Times New Roman'
            prof_run.font.size = Pt(11)

            skills_run = para.add_run(', '.join(skill_list))
            skills_run.font.name = 'Times New Roman'
            skills_run.font.size = Pt(11)
            para.paragraph_format.space_after = Pt(2)

        document.add_paragraph().paragraph_format.space_after = Pt(0)

    # ==================== WORK EXPERIENCE ====================
    if work_experiences:
        add_section_heading(document, 'EXPERIENCE')

        for work in work_experiences:
            # Company name with date range on right
            date_start = format_date(work['start_date'])
            date_end = format_date(work['end_date']) if not work.get('is_current') else "Present"
            date_range = f"{date_start} - {date_end}"

            add_entry_with_date(document, f"{work['company']}", date_range, bold_left=True)

            # Job title
            add_subtitle(document, work['title'])

            # Responsibilities as bullet points
            if work.get('responsibilities'):
                # Split responsibilities by newline or period if multiple
                responsibilities = work['responsibilities']
                # Try to split by newlines first
                resp_list = [r.strip() for r in responsibilities.split('\n') if r.strip()]

                # If no newlines, treat as single responsibility
                if len(resp_list) <= 1:
                    resp_list = [responsibilities]

                for resp in resp_list:
                    if resp:
                        add_bullet_point(document, resp)

            document.add_paragraph().paragraph_format.space_after = Pt(0)

    # ==================== PROJECTS ====================
    if projects:
        add_section_heading(document, 'PROJECTS')

        for project in projects:
            # Project title with date on right
            date_range = ""
            if project.get('start_date'):
                date_start = format_date(project['start_date'])
                date_end = format_date(project['end_date']) if project.get('end_date') else "Ongoing"
                date_range = f"{date_start} - {date_end}"

            add_entry_with_date(document, project['title'], date_range, bold_left=True)

            # Technologies
            if project.get('technologies'):
                tech_para = document.add_paragraph()
                tech_run = tech_para.add_run(f"Technologies: {project['technologies']}")
                tech_run.font.name = 'Times New Roman'
                tech_run.font.size = Pt(10)
                tech_run.font.italic = True
                tech_para.paragraph_format.space_after = Pt(2)

            # Description as bullet points
            if project.get('description'):
                desc_list = [d.strip() for d in project['description'].split('\n') if d.strip()]
                if len(desc_list) <= 1:
                    desc_list = [project['description']]

                for desc in desc_list:
                    if desc:
                        add_bullet_point(document, desc)

            # URL
            if project.get('url'):
                url_para = document.add_paragraph()
                url_run = url_para.add_run(f"Link: {project['url']}")
                url_run.font.name = 'Times New Roman'
                url_run.font.size = Pt(10)
                url_run.font.color.rgb = RGBColor(0, 0, 255)
                url_para.paragraph_format.space_after = Pt(2)

            document.add_paragraph().paragraph_format.space_after = Pt(0)

    # Save to BytesIO
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream


def generate_tailored_resume(tailored_data: dict, job_title: str) -> io.BytesIO:
    """
    Generate a tailored DOCX resume using AI-rephrased content
    tailored_data contains: user, education, original_work_experiences, original_projects, original_skills, tailored
    """

    user = tailored_data['user']
    education = tailored_data['education']
    original_work = tailored_data['original_work_experiences']
    original_projects = tailored_data['original_projects']
    original_skills = tailored_data['original_skills']
    tailored = tailored_data['tailored']

    # Create document
    document = Document()

    # Set margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # ==================== HEADER - NAME ====================
    name_para = document.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(user['name'])
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.name = 'Times New Roman'
    name_para.paragraph_format.space_after = Pt(0)

    # ==================== CONTACT INFO ====================
    contact_para = document.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = []

    if user.get('location'):
        contact_parts.append(user['location'])
    if user.get('phone'):
        contact_parts.append(user['phone'])
    if user.get('email'):
        contact_parts.append(user['email'])

    contact_run = contact_para.add_run(' | '.join(contact_parts))
    contact_run.font.size = Pt(10)
    contact_run.font.name = 'Times New Roman'
    contact_para.paragraph_format.space_after = Pt(6)

    # ==================== PROFESSIONAL SUMMARY ====================
    if tailored.get('tailored_summary'):
        add_section_heading(document, 'PROFESSIONAL SUMMARY')
        summary_para = document.add_paragraph()
        summary_run = summary_para.add_run(tailored['tailored_summary'])
        summary_run.font.name = 'Times New Roman'
        summary_run.font.size = Pt(11)
        summary_para.paragraph_format.space_after = Pt(6)

    # ==================== EDUCATION ====================
    if education:
        add_section_heading(document, 'EDUCATION')

        for edu in education:
            school_text = f"{edu['school']}"
            date_text = format_date(edu['end_date']) if edu.get('end_date') else ""
            add_entry_with_date(document, school_text, date_text, bold_left=True)

            degree_text = edu['degree']
            if edu.get('field_of_study'):
                degree_text += f", {edu['field_of_study']}"
            add_subtitle(document, degree_text)

            if edu.get('gpa'):
                gpa_para = document.add_paragraph()
                run = gpa_para.add_run(f"GPA: {edu['gpa']}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                gpa_para.paragraph_format.space_after = Pt(6)
            else:
                document.add_paragraph().paragraph_format.space_after = Pt(0)

    # ==================== SKILLS (Tailored order) ====================
    tailored_skills = tailored.get('tailored_skills', [])
    if tailored_skills:
        add_section_heading(document, 'SKILLS')
        skills_para = document.add_paragraph()
        skills_run = skills_para.add_run(', '.join(tailored_skills))
        skills_run.font.name = 'Times New Roman'
        skills_run.font.size = Pt(11)
        skills_para.paragraph_format.space_after = Pt(6)

    # ==================== WORK EXPERIENCE (Tailored) ====================
    tailored_work = tailored.get('tailored_work_experiences', [])
    if tailored_work or original_work:
        add_section_heading(document, 'EXPERIENCE')

        # Create a map of tailored work by ID
        tailored_work_map = {w.get('id'): w for w in tailored_work}

        for work in original_work:
            # Get tailored version if available
            tailored_version = tailored_work_map.get(work['id'], {})

            date_start = format_date(work['start_date'])
            date_end = format_date(work['end_date']) if not work.get('is_current') else "Present"
            date_range = f"{date_start} - {date_end}"

            company = tailored_version.get('company', work['company'])
            title = tailored_version.get('title', work['title'])

            add_entry_with_date(document, company, date_range, bold_left=True)
            add_subtitle(document, title)

            # Use tailored responsibilities if available, otherwise original
            responsibilities = tailored_version.get('responsibilities', work.get('responsibilities', ''))

            if responsibilities:
                resp_list = [r.strip() for r in responsibilities.split('\n') if r.strip()]
                if len(resp_list) <= 1 and responsibilities:
                    resp_list = [responsibilities]

                for resp in resp_list:
                    if resp:
                        add_bullet_point(document, resp)

            document.add_paragraph().paragraph_format.space_after = Pt(0)

    # ==================== PROJECTS (Tailored) ====================
    tailored_projects = tailored.get('tailored_projects', [])
    if tailored_projects or original_projects:
        add_section_heading(document, 'PROJECTS')

        # Create a map of tailored projects by ID
        tailored_proj_map = {p.get('id'): p for p in tailored_projects}

        for project in original_projects:
            tailored_version = tailored_proj_map.get(project['id'], {})

            date_range = ""
            if project.get('start_date'):
                date_start = format_date(project['start_date'])
                date_end = format_date(project['end_date']) if project.get('end_date') else "Ongoing"
                date_range = f"{date_start} - {date_end}"

            title = tailored_version.get('title', project['title'])
            add_entry_with_date(document, title, date_range, bold_left=True)

            if project.get('technologies'):
                tech_para = document.add_paragraph()
                tech_run = tech_para.add_run(f"Technologies: {project['technologies']}")
                tech_run.font.name = 'Times New Roman'
                tech_run.font.size = Pt(10)
                tech_run.font.italic = True
                tech_para.paragraph_format.space_after = Pt(2)

            # Use tailored description if available
            description = tailored_version.get('description', project.get('description', ''))

            if description:
                desc_list = [d.strip() for d in description.split('\n') if d.strip()]
                if len(desc_list) <= 1 and description:
                    desc_list = [description]

                for desc in desc_list:
                    if desc:
                        add_bullet_point(document, desc)

            if project.get('url'):
                url_para = document.add_paragraph()
                url_run = url_para.add_run(f"Link: {project['url']}")
                url_run.font.name = 'Times New Roman'
                url_run.font.size = Pt(10)
                url_run.font.color.rgb = RGBColor(0, 0, 255)
                url_para.paragraph_format.space_after = Pt(2)

            document.add_paragraph().paragraph_format.space_after = Pt(0)

    # Save to BytesIO
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream
