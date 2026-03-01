from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import psycopg2
from psycopg2.extras import RealDictCursor
from .database import get_connection
import io


FONT_NAME = 'Arial'


def format_date_month_year(date_obj):
    """Format date object to 'Mon YYYY' (e.g., 'May 2025')"""
    if date_obj:
        return date_obj.strftime("%b %Y")
    return "Present"


def format_date_range(start_date, end_date, is_current=False):
    """Format date range like 'May 2025 – Aug 2025' or 'Sep 2023 – Present'"""
    start = format_date_month_year(start_date)
    if is_current:
        end = "Present"
    else:
        end = format_date_month_year(end_date)
    return f"{start} \u2013 {end}" if start else end


def format_expected_date(end_date):
    """Format expected graduation like 'Expected 2027'"""
    if end_date:
        return f"Expected {end_date.strftime('%Y')}"
    return ""


def set_paragraph_spacing(para, before=0, after=0, line_spacing=None):
    """Set paragraph spacing"""
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)


def add_bottom_border(para):
    """Add a thin bottom border line to a paragraph"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(document, text):
    """Add a section heading with bottom border (e.g., EDUCATION, EXPERIENCE)"""
    para = document.add_paragraph()
    run = para.add_run(text.upper())
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0, 0, 0)
    add_bottom_border(para)
    set_paragraph_spacing(para, before=12, after=4)
    return para


def add_entry_line(document, left_text, right_text, left_bold=True, left_size=10, right_size=10):
    """Add a line with left-aligned text and right-aligned text using tab stops"""
    para = document.add_paragraph()

    # Set right-aligned tab stop at the right margin
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')  # ~6.5 inches in twips
    tabs.append(tab)
    pPr.append(tabs)

    left_run = para.add_run(left_text)
    left_run.font.name = FONT_NAME
    left_run.font.size = Pt(left_size)
    left_run.font.bold = left_bold
    left_run.font.color.rgb = RGBColor(0, 0, 0)

    if right_text:
        tab_run = para.add_run('\t')
        tab_run.font.size = Pt(right_size)

        right_run = para.add_run(right_text)
        right_run.font.name = FONT_NAME
        right_run.font.size = Pt(right_size)
        right_run.font.color.rgb = RGBColor(0, 0, 0)

    set_paragraph_spacing(para, before=0, after=0)
    return para


def add_subtitle_line(document, left_text, right_text=""):
    """Add an italic subtitle line (e.g., degree, job title)"""
    para = document.add_paragraph()

    if right_text:
        pPr = para._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9360')
        tabs.append(tab)
        pPr.append(tabs)

    run = para.add_run(left_text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    if right_text:
        tab_run = para.add_run('\t')
        tab_run.font.size = Pt(10)
        right_run = para.add_run(right_text)
        right_run.font.name = FONT_NAME
        right_run.font.size = Pt(10)
        right_run.font.italic = True
        right_run.font.color.rgb = RGBColor(0, 0, 0)

    set_paragraph_spacing(para, before=0, after=2)
    return para


def add_bullet_point(document, text):
    """Add a bullet point item"""
    para = document.add_paragraph()
    run = para.add_run(f"\u2022 {text}")
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.first_line_indent = Inches(-0.15)
    set_paragraph_spacing(para, before=0, after=1)
    return para


def add_skills_line(document, category, items):
    """Add a skills line like 'Languages: Python, SQL'"""
    para = document.add_paragraph()

    cat_run = para.add_run(f"{category}: ")
    cat_run.font.name = FONT_NAME
    cat_run.font.size = Pt(10)
    cat_run.font.bold = True
    cat_run.font.color.rgb = RGBColor(0, 0, 0)

    items_run = para.add_run(items)
    items_run.font.name = FONT_NAME
    items_run.font.size = Pt(10)
    items_run.font.color.rgb = RGBColor(0, 0, 0)

    set_paragraph_spacing(para, before=0, after=1)
    return para


def categorize_skills(skills):
    """Categorize skills into structured groups for Big Tech format.
    Returns list of (category, comma-separated skills) tuples."""
    language_keywords = {'python', 'java', 'javascript', 'typescript', 'c', 'c++', 'c#', 'go',
                         'rust', 'ruby', 'php', 'swift', 'kotlin', 'scala', 'r', 'matlab',
                         'sql', 'html', 'css', 'bash', 'shell', 'perl', 'lua', 'dart',
                         'objective-c', 'assembly', 'haskell', 'elixir', 'clojure'}
    backend_keywords = {'fastapi', 'django', 'flask', 'express', 'node.js', 'nodejs', 'spring',
                        'spring boot', 'nest.js', 'nestjs', 'rails', 'laravel', 'asp.net',
                        'rest api', 'rest apis', 'graphql', 'grpc', 'microservices'}
    frontend_keywords = {'react', 'vue', 'angular', 'svelte', 'next.js', 'nextjs', 'nuxt',
                         'tailwind', 'bootstrap', 'material ui', 'redux', 'jquery',
                         'html/css', 'sass', 'less', 'webpack', 'vite'}
    database_keywords = {'postgresql', 'postgres', 'mysql', 'mongodb', 'redis', 'sqlite',
                         'dynamodb', 'cassandra', 'elasticsearch', 'firebase', 'supabase',
                         'oracle', 'sql server', 'mariadb', 'neo4j'}
    devops_keywords = {'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'terraform', 'ansible',
                       'jenkins', 'ci/cd', 'nginx', 'linux', 'heroku', 'vercel', 'netlify',
                       'railway', 'render', 'cloudflare'}
    tools_keywords = {'git', 'github', 'gitlab', 'bitbucket', 'jira', 'figma', 'postman',
                      'vscode', 'vim', 'intellij', 'jupyter', 'slack', 'notion'}

    categories = {
        'Languages': [],
        'Backend': [],
        'Frontend': [],
        'Database': [],
        'DevOps': [],
        'Tools': [],
    }

    for skill in skills:
        name = skill['skill_name']
        name_lower = name.lower().strip()

        if name_lower in language_keywords:
            categories['Languages'].append(name)
        elif name_lower in backend_keywords:
            categories['Backend'].append(name)
        elif name_lower in frontend_keywords:
            categories['Frontend'].append(name)
        elif name_lower in database_keywords:
            categories['Database'].append(name)
        elif name_lower in devops_keywords:
            categories['DevOps'].append(name)
        elif name_lower in tools_keywords:
            categories['Tools'].append(name)
        else:
            # Try partial matching for multi-word skills
            matched = False
            for kw in language_keywords:
                if kw in name_lower or name_lower in kw:
                    categories['Languages'].append(name)
                    matched = True
                    break
            if not matched:
                for kw in backend_keywords:
                    if kw in name_lower or name_lower in kw:
                        categories['Backend'].append(name)
                        matched = True
                        break
            if not matched:
                for kw in frontend_keywords:
                    if kw in name_lower or name_lower in kw:
                        categories['Frontend'].append(name)
                        matched = True
                        break
            if not matched:
                for kw in database_keywords:
                    if kw in name_lower or name_lower in kw:
                        categories['Database'].append(name)
                        matched = True
                        break
            if not matched:
                for kw in devops_keywords:
                    if kw in name_lower or name_lower in kw:
                        categories['DevOps'].append(name)
                        matched = True
                        break
            if not matched:
                categories['Tools'].append(name)

    # Return only non-empty categories
    return [(cat, ', '.join(items)) for cat, items in categories.items() if items]


def build_header(document, user):
    """Build the ATS-safe header: name, contact line, links line"""
    # Name - centered, bold, larger font
    name_para = document.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(user['name'].upper())
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.name = FONT_NAME
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    set_paragraph_spacing(name_para, before=0, after=2)

    # Contact line: Location | Phone | Email
    contact_parts = []
    if user.get('location'):
        contact_parts.append(user['location'])
    if user.get('phone'):
        contact_parts.append(user['phone'])
    if user.get('email'):
        contact_parts.append(user['email'])

    if contact_parts:
        contact_para = document.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(' | '.join(contact_parts))
        contact_run.font.name = FONT_NAME
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(0, 0, 0)
        set_paragraph_spacing(contact_para, before=0, after=2)

    # Links line: GitHub | LinkedIn
    link_parts = []
    if user.get('github'):
        link_parts.append(user['github'])
    if user.get('linkedin'):
        link_parts.append(user['linkedin'])

    if link_parts:
        links_para = document.add_paragraph()
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_run = links_para.add_run(' | '.join(link_parts))
        links_run.font.name = FONT_NAME
        links_run.font.size = Pt(10)
        links_run.font.color.rgb = RGBColor(0, 0, 0)
        set_paragraph_spacing(links_para, before=0, after=4)


def build_education_section(document, education):
    """Build the Education section"""
    if not education:
        return

    add_section_heading(document, 'EDUCATION')

    for edu in education:
        # School name — right-aligned location or date
        expected = format_expected_date(edu.get('end_date'))
        date_range = format_date_range(edu.get('start_date'), edu.get('end_date'))

        add_entry_line(document, edu['school'], expected or date_range, left_bold=True)

        # Degree line
        degree_text = edu['degree']
        if edu.get('field_of_study'):
            degree_text += f", {edu['field_of_study']}"
        if expected:
            degree_text += f" \u2014 {expected}"

        add_subtitle_line(document, degree_text)

        # GPA
        if edu.get('gpa'):
            add_bullet_point(document, f"GPA: {edu['gpa']}")

        # Description as relevant coursework or achievements
        if edu.get('description'):
            lines = [l.strip() for l in edu['description'].split('\n') if l.strip()]
            for line in lines:
                add_bullet_point(document, line)


def build_experience_section(document, work_experiences, heading='EXPERIENCE'):
    """Build the Experience section in Big Tech format"""
    if not work_experiences:
        return

    add_section_heading(document, heading)

    for work in work_experiences:
        # COMPANY — Title    right-aligned dates
        company_title = f"{work['company']} \u2014 {work['title']}"
        date_range = format_date_range(work.get('start_date'), work.get('end_date'), work.get('is_current'))

        add_entry_line(document, company_title, date_range, left_bold=True)

        # Bullet points for responsibilities
        if work.get('responsibilities'):
            resp_list = [r.strip() for r in work['responsibilities'].split('\n') if r.strip()]
            for resp in resp_list:
                # Strip leading bullet/dash if already present
                resp = resp.lstrip('\u2022-* ').strip()
                if resp:
                    add_bullet_point(document, resp)


def build_projects_section(document, projects):
    """Build the Projects section in Big Tech format"""
    if not projects:
        return

    add_section_heading(document, 'PROJECTS')

    for project in projects:
        # Title | Technologies
        title_text = project['title']
        if project.get('technologies'):
            title_text += f" | {project['technologies']}"

        add_entry_line(document, title_text, "", left_bold=True)

        # Description bullet points
        if project.get('description'):
            desc_list = [d.strip() for d in project['description'].split('\n') if d.strip()]
            for desc in desc_list:
                desc = desc.lstrip('\u2022-* ').strip()
                if desc:
                    add_bullet_point(document, desc)


def build_skills_section(document, skills):
    """Build the Skills section in structured category format"""
    if not skills:
        return

    add_section_heading(document, 'SKILLS')

    categorized = categorize_skills(skills)
    for category, items in categorized:
        add_skills_line(document, category, items)


def create_base_document():
    """Create a new DOCX document with standard settings"""
    document = Document()

    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = document.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    # Remove default paragraph spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    return document


def generate_resume(user_id: int) -> io.BytesIO:
    """
    Generate a clean, ATS-safe Big Tech resume as DOCX.
    No icons, no symbols, no fluff.
    """
    conn = get_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)

    try:
        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        if not user:
            raise ValueError("User not found")

        cursor.execute("""
            SELECT * FROM work_experiences
            WHERE user_id = %s
            ORDER BY is_current DESC, start_date DESC
        """, (user_id,))
        work_experiences = cursor.fetchall()

        cursor.execute("""
            SELECT * FROM education
            WHERE user_id = %s
            ORDER BY end_date DESC NULLS FIRST
        """, (user_id,))
        education = cursor.fetchall()

        cursor.execute("""
            SELECT * FROM skills
            WHERE user_id = %s
            ORDER BY skill_name ASC
        """, (user_id,))
        skills = cursor.fetchall()

        cursor.execute("""
            SELECT * FROM projects
            WHERE user_id = %s
            ORDER BY start_date DESC NULLS LAST
        """, (user_id,))
        projects = cursor.fetchall()
    finally:
        cursor.close()
        conn.close()

    document = create_base_document()

    # HEADER
    build_header(document, user)

    # EDUCATION (comes before experience for students)
    build_education_section(document, education)

    # EXPERIENCE
    build_experience_section(document, work_experiences)

    # PROJECTS
    build_projects_section(document, projects)

    # SKILLS
    build_skills_section(document, skills)

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


def generate_tailored_resume(tailored_data: dict, job_title: str) -> io.BytesIO:
    """
    Generate a tailored Big Tech resume using AI-rephrased content.
    Same clean format, but with content optimized for the target role.
    """
    user = tailored_data['user']
    education = tailored_data['education']
    original_work = tailored_data['original_work_experiences']
    original_projects = tailored_data['original_projects']
    original_skills = tailored_data['original_skills']
    tailored = tailored_data['tailored']

    document = create_base_document()

    # HEADER
    build_header(document, user)

    # EDUCATION
    build_education_section(document, education)

    # EXPERIENCE (Tailored)
    tailored_work = tailored.get('tailored_work_experiences', [])
    tailored_work_map = {w.get('id'): w for w in tailored_work}

    merged_work = []
    for work in original_work:
        tailored_version = tailored_work_map.get(work['id'], {})
        merged = {
            'company': tailored_version.get('company', work['company']),
            'title': tailored_version.get('title', work['title']),
            'start_date': work.get('start_date'),
            'end_date': work.get('end_date'),
            'is_current': work.get('is_current'),
            'responsibilities': tailored_version.get('responsibilities', work.get('responsibilities', '')),
        }
        merged_work.append(merged)

    build_experience_section(document, merged_work)

    # PROJECTS (Tailored)
    tailored_projects = tailored.get('tailored_projects', [])
    tailored_proj_map = {p.get('id'): p for p in tailored_projects}

    merged_projects = []
    for project in original_projects:
        tailored_version = tailored_proj_map.get(project['id'], {})
        merged = {
            'title': tailored_version.get('title', project['title']),
            'technologies': project.get('technologies', ''),
            'description': tailored_version.get('description', project.get('description', '')),
            'start_date': project.get('start_date'),
            'end_date': project.get('end_date'),
        }
        merged_projects.append(merged)

    build_projects_section(document, merged_projects)

    # SKILLS (Tailored order)
    tailored_skills = tailored.get('tailored_skills', [])
    if tailored_skills:
        # Convert string list to skill dicts for categorize_skills
        skill_dicts = [{'skill_name': s} for s in tailored_skills]
        add_section_heading(document, 'SKILLS')
        categorized = categorize_skills(skill_dicts)
        for category, items in categorized:
            add_skills_line(document, category, items)
    else:
        build_skills_section(document, original_skills)

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream
